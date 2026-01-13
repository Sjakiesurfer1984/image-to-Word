# =====================================================================
# Photo appendix builder
# =====================================================================
# Purpose:
# - Optionally read an Excel list of required photo base names (Column A).
# - Let the user choose a source folder of images.
# - Let the user choose ordering: Excel order or numeric ascending.
# - Auto-tune JPEG compression so the Word document targets <= 20 MB.
# - Show a preview of the compressed image for user approval.
# - Optionally write compressed copies into a safe subfolder.
# - Generate a Word document in landscape layout with a 2x3 photo grid per page.
#
# Style choices:
# - Australian English spelling in comments and messages.
# - Type hints everywhere.
# - Comments focus on why something is done, not only what.
# =====================================================================

import io
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from pathlib import Path
import shutil
import tempfile
from typing import Dict, List, Optional, Sequence, Tuple

from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook
import re


# =====================================================================
# Constants and configuration
# =====================================================================

TARGET_DOCX_SIZE_MB: float = 20.0
TARGET_DOCX_SIZE_BYTES: int = int(TARGET_DOCX_SIZE_MB * 1024 * 1024)

# DOCX contains images plus XML and relationship metadata. This is a conservative allowance.
DOCX_OVERHEAD_BYTES: int = 1_500_000

# Guard rails so the auto-tuner does not destroy image usefulness.
QUALITY_MIN: int = 30
QUALITY_MAX: int = 90
MAX_DIM_MIN: int = 800
MAX_DIM_MAX: int = 2400

# Auto-tune step sizes for predictable behaviour.
QUALITY_STEP: int = 5
MAX_DIM_STEP: int = 200

# Sample size for estimating output size quickly.
ESTIMATE_SAMPLE_MAX: int = 6

# Supported image extensions.
SUPPORTED_SUFFIXES: set[str] = {".jpg", ".jpeg", ".png", ".bmp"}


# =====================================================================
# Module 1: Excel input
# =====================================================================

def get_photo_list_from_excel() -> List[str]:
    """
    Prompts the user for an Excel file and extracts values from the first column
    of the active worksheet. Each non-empty cell is treated as a required photo
    base name (matching Path.stem) to filter the source images.

    Why read_only=True:
    - Avoids loading the entire workbook into memory.
    - More robust for simple read tasks and large files.
    """
    print("\nPlease select the Excel file containing the list of required photo filenames.")
    file_path_str: str = filedialog.askopenfilename(
        title="Select Excel Photo List (.xlsx)",
        filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
    )

    if not file_path_str:
        print("No Excel file selected.")
        return []

    file_path: Path = Path(file_path_str)

    try:
        workbook = load_workbook(filename=file_path, read_only=True)
        sheet = workbook.active

        photo_list: List[str] = []
        for i in range(1, sheet.max_row + 1):
            cell = sheet.cell(row=i, column=1)
            if cell.value is None:
                continue
            value: str = str(cell.value).strip()
            if value:
                photo_list.append(value)

        print(f"Successfully read {len(photo_list)} entries from Excel file.")
        return photo_list

    except Exception as e:
        print(f"Error reading Excel file: {e}")
        messagebox.showerror("File Error", f"Could not read the selected Excel file.\n\nError: {e}")
        return []


def natural_sort_key_string(value: str) -> List[object]:
    """
    Natural sorting for stems so '10' does not come before '2'.
    """
    cleaned: str = str(value).strip()
    parts: List[str] = re.split(r"(\d+)", cleaned)
    return [int(p) if p.isdigit() else p for p in parts]


# =====================================================================
# Module 2: Image processing for Word embedding and compression
# =====================================================================

def _apply_exif_orientation(img: Image.Image) -> Image.Image:
    """
    Applies EXIF orientation correction if present.

    Why:
    - Many cameras store orientation in metadata, not pixels.
    - Fixing orientation ensures the Word output matches what people expect.
    """
    try:
        exif = getattr(img, "_getexif", lambda: None)()
        if not exif:
            return img

        orientation_tag: int = 274
        orientation = exif.get(orientation_tag)

        if orientation == 3:
            return img.rotate(180, expand=True)
        if orientation == 6:
            return img.rotate(270, expand=True)
        if orientation == 8:
            return img.rotate(90, expand=True)

        return img
    except Exception:
        return img


def compress_image_to_jpeg_bytes(
    img_path: Path,
    quality: int,
    max_dim: int,
) -> bytes:
    """
    Creates a compressed JPEG representation in memory.

    Why bytes:
    - Enables fast estimation of expected output size without writing files.
    """
    img: Image.Image = Image.open(img_path)
    img = _apply_exif_orientation(img)

    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")

    img.thumbnail((max_dim, max_dim), Image.Resampling.LANCZOS)

    bio = io.BytesIO()
    img.save(
        bio,
        format="JPEG",
        quality=quality,
        optimize=True,
        progressive=True,
    )
    return bio.getvalue()


def write_prepared_jpeg_for_word(
    img_path: Path,
    out_dir: Path,
    quality: int,
    max_dim: int,
) -> Path:
    """
    Writes a prepared JPEG file suitable for Word embedding.

    Why write files:
    - python-docx expects a stable on-disk path for add_picture.
    - Keeping all processed assets together makes debugging simpler.
    """
    jpg_bytes: bytes = compress_image_to_jpeg_bytes(img_path, quality=quality, max_dim=max_dim)
    out_path: Path = out_dir / f"{img_path.stem}.jpg"
    out_path.write_bytes(jpg_bytes)
    return out_path


def estimate_docx_size_bytes(
    image_paths: Sequence[Path],
    quality: int,
    max_dim: int,
    overhead_bytes: int = DOCX_OVERHEAD_BYTES,
    sample_max: int = ESTIMATE_SAMPLE_MAX,
) -> int:
    """
    Estimates the final Word document size by compressing a sample set in memory.

    Strategy:
    - Use the largest source images because they are most likely to dominate output size.
    - Compute average compressed size per image based on sample.
    - Multiply by total number of images and add DOCX overhead.

    This is an estimate, not an exact value, but it is close enough to guide tuning.
    """
    if not image_paths:
        return overhead_bytes

    sorted_by_size: List[Path] = sorted(image_paths, key=lambda p: p.stat().st_size, reverse=True)
    sample: List[Path] = sorted_by_size[:min(sample_max, len(sorted_by_size))]

    compressed_sizes: List[int] = []
    for p in sample:
        try:
            b = compress_image_to_jpeg_bytes(p, quality=quality, max_dim=max_dim)
            compressed_sizes.append(len(b))
        except Exception:
            continue

    if not compressed_sizes:
        return overhead_bytes + 10_000_000

    avg_size: float = sum(compressed_sizes) / float(len(compressed_sizes))
    estimate: int = int(overhead_bytes + avg_size * len(image_paths))
    return estimate


def auto_tune_compression(
    image_paths: Sequence[Path],
    target_bytes: int,
    start_quality: int,
    start_max_dim: int,
) -> Tuple[int, int, int]:
    """
    Finds a (quality, max_dim) pair that estimates under the target size.

    Tuning rules:
    - Reduce quality first, because reducing pixel count can degrade readability faster.
    - If quality hits a floor, reduce max_dim in steps.
    - Stops when it meets target or reaches minimum bounds.

    Returns:
    - chosen_quality
    - chosen_max_dim
    - estimated_bytes
    """
    quality: int = max(QUALITY_MIN, min(QUALITY_MAX, int(start_quality)))
    max_dim: int = max(MAX_DIM_MIN, min(MAX_DIM_MAX, int(start_max_dim)))

    while True:
        est: int = estimate_docx_size_bytes(image_paths, quality=quality, max_dim=max_dim)

        if est <= target_bytes:
            return quality, max_dim, est

        if quality - QUALITY_STEP >= QUALITY_MIN:
            quality -= QUALITY_STEP
            continue

        if max_dim - MAX_DIM_STEP >= MAX_DIM_MIN:
            max_dim -= MAX_DIM_STEP
            continue

        # Cannot reduce further without breaking minimum limits.
        return quality, max_dim, est


def show_compression_preview_dialog(
    parent: tk.Tk,
    sample_image_path: Path,
    quality: int,
    max_dim: int,
    estimated_bytes: int,
    target_bytes: int,
) -> bool:
    """
    Shows a preview of the compressed sample image and asks the user to proceed.

    Returns:
    - True if user accepts
    - False if user cancels
    """
    try:
        jpg_bytes: bytes = compress_image_to_jpeg_bytes(sample_image_path, quality=quality, max_dim=max_dim)
        img = Image.open(io.BytesIO(jpg_bytes))
    except Exception as e:
        messagebox.showerror("Preview error", f"Could not create a preview image.\n\nError: {e}")
        return False

    top = tk.Toplevel(parent)
    top.title("Compression preview")

    info_text: str = (
        f"Sample photo: {sample_image_path.name}\n"
        f"JPEG quality: {quality}\n"
        f"Max dimension: {max_dim}px\n\n"
        f"Estimated Word doc size: {estimated_bytes / (1024 * 1024):.2f} MB\n"
        f"Target size: {target_bytes / (1024 * 1024):.2f} MB\n\n"
        "Proceed with these settings?"
    )

    lbl = tk.Label(top, text=info_text, justify="left")
    lbl.pack(padx=12, pady=10)

    # Fit the preview to a reasonable UI size.
    preview_max_ui: int = 900
    img.thumbnail((preview_max_ui, preview_max_ui), Image.Resampling.LANCZOS)

    tk_img = ImageTk.PhotoImage(img)
    img_label = tk.Label(top, image=tk_img)
    img_label.image = tk_img
    img_label.pack(padx=12, pady=10)

    result: Dict[str, bool] = {"ok": False}

    def on_ok() -> None:
        result["ok"] = True
        top.destroy()

    def on_cancel() -> None:
        result["ok"] = False
        top.destroy()

    btn_frame = tk.Frame(top)
    btn_frame.pack(pady=12)

    ok_btn = tk.Button(btn_frame, text="Proceed", width=14, command=on_ok)
    ok_btn.pack(side="left", padx=8)

    cancel_btn = tk.Button(btn_frame, text="Cancel", width=14, command=on_cancel)
    cancel_btn.pack(side="left", padx=8)

    top.grab_set()
    parent.wait_window(top)
    return result["ok"]


# =====================================================================
# Module 3: Word document generation using prepared JPEGs
# =====================================================================

def create_photo_document_from_prepared_jpegs(
    prepared_jpeg_paths: List[Path],
    save_directory: Path,
) -> Path:
    """
    Creates a landscape Word document with a 2x3 image grid per page.

    This function assumes images are already pre-compressed and resized.
    """
    if not prepared_jpeg_paths:
        raise ValueError("No prepared images were provided for document creation.")

    doc: Document = Document()
    section = doc.sections[-1]
    section.orientation = 1
    section.page_width, section.page_height = section.page_height, section.page_width

    images_per_page: int = 6
    cols_per_row: int = 3
    img_width: Inches = Inches(2.5)

    for start in range(0, len(prepared_jpeg_paths), images_per_page):
        page_imgs: List[Path] = prepared_jpeg_paths[start:start + images_per_page]

        for row_index in range(2):
            table = doc.add_table(rows=1, cols=cols_per_row)
            row_cells = table.rows[0].cells

            for col_index in range(cols_per_row):
                img_index: int = row_index * cols_per_row + col_index
                if img_index >= len(page_imgs):
                    continue

                jpg_path: Path = page_imgs[img_index]

                run = row_cells[col_index].paragraphs[0].add_run()
                run.add_picture(str(jpg_path), width=img_width)

                caption = row_cells[col_index].add_paragraph(jpg_path.stem)
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption.runs[0].font.name = "Aptos"
                caption.runs[0].font.size = Pt(10)

        if start + images_per_page < len(prepared_jpeg_paths):
            doc.add_page_break()

    output_path: Path = save_directory / "Photo_Appendix.docx"
    doc.save(output_path)
    return output_path


def save_compressed_images_from_prepared(
    prepared_jpeg_paths: List[Path],
    source_directory: Path,
) -> Path:
    """
    Writes compressed copies into 'COMPRESSED_IMAGES' in the source folder.

    These are already compressed. This function copies them so colleagues can
    reuse the smaller images outside of Word if they want.
    """
    output_dir: Path = source_directory / "COMPRESSED_IMAGES"
    output_dir.mkdir(exist_ok=True)

    for p in prepared_jpeg_paths:
        shutil.copy2(p, output_dir / p.name)

    return output_dir


# =====================================================================
# Controller: Main orchestration logic
# =====================================================================

def main_orchestrator() -> None:
    root = tk.Tk()
    root.withdraw()

    source_dir_str: str = filedialog.askdirectory(title="Select the folder of photos for the report")
    if not source_dir_str:
        print("No source folder selected. Exiting.")
        return

    source_dir: Path = Path(source_dir_str)

    all_images: List[Path] = [
        p for p in source_dir.glob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
    ]

    if not all_images:
        messagebox.showwarning("No photos", "No supported image files were found in the selected folder.")
        return

    photos_for_report: List[Path] = []

    use_excel: bool = messagebox.askyesno(
        "Filter photos?",
        "Do you want to use an Excel file to select specific photos for the report?"
    )

    if use_excel:
        required_photo_list: List[str] = get_photo_list_from_excel()
        if not required_photo_list:
            print("No photo list retrieved from Excel, or dialog was cancelled. Aborting process.")
            return

        sort_numeric: bool = messagebox.askyesno(
            "Photo ordering",
            "How should photos be ordered in the Word document?\n\n"
            "Yes = sort numerically (ascending)\n"
            "No = use the native Excel order"
        )

        ordered_stems: List[str] = (
            sorted(required_photo_list, key=natural_sort_key_string) if sort_numeric else required_photo_list
        )

        image_by_stem: Dict[str, Path] = {p.stem: p for p in all_images}
        missing_stems: List[str] = []

        for stem in ordered_stems:
            stem_clean: str = str(stem).strip()
            match: Optional[Path] = image_by_stem.get(stem_clean)
            if match is not None:
                photos_for_report.append(match)
            else:
                missing_stems.append(stem_clean)

        if missing_stems:
            messagebox.showwarning(
                "Some photos not found",
                f"{len(missing_stems)} photo(s) listed in Excel were not found in the selected folder.\n\n"
                "The document will be created with the photos that were found."
            )

    else:
        # Folder mode without Excel: natural ordering by stem is usually what people expect.
        photos_for_report = sorted(all_images, key=lambda p: natural_sort_key_string(p.stem))

    if not photos_for_report:
        messagebox.showwarning("No photos", "No photos were selected or matched for the report.")
        return

    print(f"\nSelected {len(photos_for_report)} photos for the report.")

    # Ask whether to auto-tune for a 20 MB cap.
    auto_tune: bool = messagebox.askyesno(
        "Word doc size control",
        f"Do you want to auto-tune compression so the Word document targets <= {TARGET_DOCX_SIZE_MB:.0f} MB?"
    )

    start_quality: int = 50
    start_max_dim: int = 1400

    if not auto_tune:
        q: Optional[int] = simpledialog.askinteger(
            "JPEG quality",
            "Enter JPEG quality (1-100). Lower means smaller file size:",
            initialvalue=start_quality, minvalue=1, maxvalue=100
        )
        d: Optional[int] = simpledialog.askinteger(
            "Max dimension",
            "Enter maximum pixel dimension for the long edge (e.g. 1200, 1400, 1600):",
            initialvalue=start_max_dim, minvalue=MAX_DIM_MIN, maxvalue=MAX_DIM_MAX
        )

        chosen_quality: int = start_quality if q is None else int(q)
        chosen_max_dim: int = start_max_dim if d is None else int(d)
        estimated_bytes: int = estimate_docx_size_bytes(photos_for_report, chosen_quality, chosen_max_dim)
    else:
        chosen_quality, chosen_max_dim, estimated_bytes = auto_tune_compression(
            photos_for_report,
            target_bytes=TARGET_DOCX_SIZE_BYTES,
            start_quality=start_quality,
            start_max_dim=start_max_dim
        )

    # Use a representative sample image for preview.
    # Picking the largest file is conservative, as it is likely to show artefacts sooner.
    sample_for_preview: Path = max(photos_for_report, key=lambda p: p.stat().st_size)

    accepted: bool = show_compression_preview_dialog(
        parent=root,
        sample_image_path=sample_for_preview,
        quality=chosen_quality,
        max_dim=chosen_max_dim,
        estimated_bytes=estimated_bytes,
        target_bytes=TARGET_DOCX_SIZE_BYTES,
    )

    if not accepted:
        print("User cancelled after preview. Exiting without generating outputs.")
        return

    also_save_compressed: bool = messagebox.askyesno(
        "Compressed copies",
        "Do you also want to save compressed copies of the selected photos into 'COMPRESSED_IMAGES'?"
    )

    # Prepare all images into a temp directory for embedding.
    temp_dir: Path = Path(tempfile.mkdtemp())
    print(f"Created temporary directory for prepared JPEGs: {temp_dir}")

    prepared: List[Path] = []
    try:
        for img_path in photos_for_report:
            try:
                prepared_path: Path = write_prepared_jpeg_for_word(
                    img_path=img_path,
                    out_dir=temp_dir,
                    quality=chosen_quality,
                    max_dim=chosen_max_dim
                )
                prepared.append(prepared_path)
            except Exception as e:
                print(f"Skipping {img_path.name} due to processing error: {e}")

        if not prepared:
            messagebox.showerror("Error", "No images could be processed. No outputs were produced.")
            return

        # Create the Word document from prepared JPEGs.
        docx_path: Path = create_photo_document_from_prepared_jpegs(prepared, save_directory=source_dir)

        # Optionally copy compressed images out for reuse.
        if also_save_compressed:
            out_dir: Path = save_compressed_images_from_prepared(prepared, source_directory=source_dir)
            print(f"Compressed copies saved to: {out_dir}")

        # Report actual file size to the user.
        actual_size_mb: float = docx_path.stat().st_size / (1024 * 1024)

        messagebox.showinfo(
            "Success",
            f"Word document created:\n{docx_path}\n\n"
            f"Actual Word doc size: {actual_size_mb:.2f} MB\n"
            f"Compression settings:\n"
            f"- JPEG quality: {chosen_quality}\n"
            f"- Max dimension: {chosen_max_dim}px"
        )

        print(f"\nWord document saved to: {docx_path}")
        print(f"Actual Word doc size: {actual_size_mb:.2f} MB")

    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)
        print("Cleaned up temporary directory.")


if __name__ == "__main__":
    main_orchestrator()





# # =====================================================================
# # Photo appendix builder
# # =====================================================================
# # Purpose:
# # - Read an optional Excel list of required photo base names.
# # - Let the user choose a source folder of images.
# # - Optionally compress copies of the selected images in a safe subfolder.
# # - Generate a Word document in landscape layout with a 2x3 photo grid per page.
# #
# # Teaching notes:
# # - GUI interactions use Tkinter's modal dialogs (no full window UI).
# # - Images are handled with Pillow (rotation by EXIF, RGB conversion).
# # - Word doc generation uses python-docx tables to implement the grid layout.
# # - Excel reading uses openpyxl in read-only mode for robustness.
# #
# # Style choices for you:
# # - Australian English spelling in comments and messages.
# # - No ordered-number comments inside code blocks.
# # - Type hints everywhere to make intent explicit.
# # - Comments focus on why something is done, not only what.
# # =====================================================================

# # --- Module imports ---

# import tkinter as tk
# from tkinter import filedialog, messagebox, simpledialog
# from pathlib import Path
# import shutil
# import tempfile
# from typing import Dict, List, Optional, Iterable

# from PIL import Image
# from docx import Document
# from docx.shared import Inches, Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from openpyxl import load_workbook
# import re


# # =====================================================================
# # Module 1: Data input and image processing
# # =====================================================================

# def get_photo_list_from_excel() -> List[str]:
#     """
#     Prompts the user for an Excel file and extracts values from the first column
#     of the active worksheet. Each non-empty cell is treated as a required photo
#     base name (matching Path.stem) to filter the source images.

#     Why this design:
#     - Using read_only=True avoids loading the entire workbook into memory,
#       which helps on large files and is safer for simple read tasks.
#     - Reading cells row-by-row avoids edge cases with high-level slicing.
#     """
#     print("\nPlease select the Excel file containing the list of required photo filenames.")
#     file_path_str: str = filedialog.askopenfilename(
#         title="Select Excel Photo List (.xlsx)",
#         filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
#     )

#     if not file_path_str:
#         print("No Excel file selected.")
#         return []

#     file_path: Path = Path(file_path_str)

#     try:
#         workbook = load_workbook(filename=file_path, read_only=True)
#         sheet = workbook.active

#         photo_list: List[str] = []
#         # Iterating from 1..max_row matches openpyxl’s 1-based indexing.
#         for i in range(1, sheet.max_row + 1):
#             cell = sheet.cell(row=i, column=1)
#             if cell.value is not None:
#                 # Casting to str ensures uniform matching against Path.stem later.
#                 value: str = str(cell.value).strip()
#                 if value:
#                     photo_list.append(value)

#         print(f"Successfully read {len(photo_list)} entries from Excel file.")
#         return photo_list

#     except Exception as e:
#         # User feedback is important for diagnosing common issues such as locked files,
#         # sheet protection, or the wrong file format.
#         print(f"Error reading Excel file: {e}")
#         messagebox.showerror("File Error", f"Could not read the selected Excel file.\n\nError: {e}")
#         return []


# def _get_compressed_image(img_path: Path) -> Image.Image:
#     """
#     Loads an image with Pillow and ensures it is in a JPEG-compatible mode.

#     Why convert modes:
#     - JPEG does not support alpha channels or palette modes.
#     - Converting to RGB avoids runtime errors when saving as JPEG.
#     """
#     img: Image.Image = Image.open(img_path)
#     if img.mode in ("RGBA", "P"):
#         img = img.convert("RGB")
#     return img


# def save_compressed_images(image_paths: List[Path], source_directory: Path, quality: int = 75) -> None:
#     """
#     Writes compressed copies of the given images into 'COMPRESSED_IMAGES' in the source folder.

#     Non-destructive by design:
#     - Originals are never overwritten.
#     - Copies go into a dedicated subfolder for easy review and rollback.

#     About JPEG quality:
#     - Must be in the 1-100 range.
#     - 75 is a reasonable default that saves space while maintaining visible quality.
#     """
#     if not image_paths:
#         print("No images were provided for compression.")
#         return

#     output_dir: Path = source_directory / "COMPRESSED_IMAGES"
#     output_dir.mkdir(exist_ok=True)

#     print(f"\n--- Starting image compression (saving to '{output_dir.name}') ---")
#     compressed_count: int = 0
#     skipped_count: int = 0

#     for img_path in image_paths:
#         try:
#             # File size fetch allows a simple before/after compression report.
#             original_size_kb: float = img_path.stat().st_size / 1024.0

#             # Preparing the image ensures valid JPEG mode.
#             compressed_img: Image.Image = _get_compressed_image(img_path)
#             save_path: Path = output_dir / img_path.name

#             # optimise=True lets Pillow attempt Huffman-table and DCT optimisations.
#             compressed_img.save(save_path, format="JPEG", quality=quality, optimize=True)

#             new_size_kb: float = save_path.stat().st_size / 1024.0
#             reduction_kb: float = max(0.0, original_size_kb - new_size_kb)
#             reduction_pct: float = (reduction_kb / original_size_kb * 100.0) if original_size_kb > 0 else 0.0

#             print(f"-> Saved compressed: {img_path.name} ({original_size_kb:.1f} KB -> {new_size_kb:.1f} KB, {reduction_pct:.1f}% smaller)")
#             compressed_count += 1

#         except Exception as e:
#             # If a single file fails (corrupt image, permission issue), we log and continue.
#             print(f"-> Error compressing {img_path.name}: {e}. Skipping.")
#             skipped_count += 1

#     print("\n--- Compression summary ---")
#     print(f"Successfully saved {compressed_count} compressed file(s) to '{output_dir.name}'.")
#     if skipped_count > 0:
#         print(f"Skipped or failed: {skipped_count} file(s).")
#     print("---------------------------")


# def _process_image_for_doc(img_path: Path, temp_dir: Path) -> Path:
#     """
#     Produces a temporary JPEG tailored for Word insertion.

#     Operations:
#     - Opens the image.
#     - Applies EXIF-based orientation fix if present.
#     - Converts to RGB if required.
#     - Saves to a temp JPG (quality ~85) and returns that temp path.

#     Why use temp files:
#     - Word embedding expects a stable on-disk file path.
#     - Avoids any mutation of original images.
#     """
#     img: Image.Image = Image.open(img_path)

#     # EXIF orientation handling: cameras often store rotation in metadata
#     # rather than rotating the pixel matrix. This aligns the visual result.
#     try:
#         exif = getattr(img, "_getexif", lambda: None)()
#         if exif:
#             orientation_tag: int = 274  # standard orientation key
#             orientation = exif.get(orientation_tag)
#             if orientation == 3:
#                 img = img.rotate(180, expand=True)
#             elif orientation == 6:
#                 img = img.rotate(270, expand=True)
#             elif orientation == 8:
#                 img = img.rotate(90, expand=True)
#     except Exception:
#         # If EXIF is malformed or absent, we simply proceed without rotation.
#         pass

#     if img.mode in ("RGBA", "P"):
#         img = img.convert("RGB")

#     temp_jpg_path: Path = temp_dir / f"{img_path.stem}_processed.jpg"
#     img.save(temp_jpg_path, format="JPEG", quality=85)
#     return temp_jpg_path


# def natural_sort_key(path: Path) -> List[object]:
#     """
#     Returns a list that mixes strings and integers to achieve natural (human-like) sorting.
#     Example:
#         DSCN3801 -> ['DSCN', 3801]
#         10 -> [10]
#     """
#     parts: List[str] = re.split(r"(\d+)", path.stem)
#     return [int(p) if p.isdigit() else p for p in parts]


# def natural_sort_key_string(value: str) -> List[object]:
#     """
#     Natural sort key for stems read from Excel.

#     Why this exists:
#     - Excel values may be numeric strings (e.g., '9091') or mixed stems.
#     - A consistent key avoids lexicographic ordering issues (e.g., 10 before 2).
#     """
#     cleaned: str = str(value).strip()
#     parts: List[str] = re.split(r"(\d+)", cleaned)
#     return [int(p) if p.isdigit() else p for p in parts]


# def create_photo_document(image_paths: List[Path], save_directory: Path, sort_images: bool = True) -> None:
#     """
#     Creates a landscape Word document with a 2x3 image grid per page.
#     Each image is centred in its cell and labelled with its base filename.

#     Layout decisions:
#     - Using tables provides a robust grid in Word, which keeps images aligned.
#     - Fixed image width lets Word keep aspect ratio and scale height automatically.
#     - Page breaks are inserted after each complete page batch, except the last.
#     """
#     if not image_paths:
#         messagebox.showinfo("No Images Found", "No images were selected or found to create the document.")
#         return

#     doc: Document = Document()
#     section = doc.sections[-1]

#     # Orientation value 1 corresponds to landscape for python-docx.
#     # Swapping width and height is necessary to apply the change cleanly.
#     section.orientation = 1
#     section.page_width, section.page_height = section.page_height, section.page_width

#     images_per_page: int = 6
#     cols_per_row: int = 3
#     img_width: Inches = Inches(2.5)

#     # A temporary folder is used to store EXIF-fixed JPEGs for embedding.
#     temp_dir: Path = Path(tempfile.mkdtemp())
#     print(f"Created temporary directory for image processing: {temp_dir}")

#     try:
#         # Sorting is useful when using all photos from a folder.
#         # When the user supplies an Excel list, the order may be intentional, so we allow sorting to be disabled.
#         if sort_images:
#             image_paths.sort(key=natural_sort_key)

#         for start in range(0, len(image_paths), images_per_page):
#             page_imgs: List[Path] = image_paths[start:start + images_per_page]

#             for row_index in range(2):
#                 table = doc.add_table(rows=1, cols=cols_per_row)
#                 row_cells = table.rows[0].cells

#                 for col_index in range(cols_per_row):
#                     img_index: int = row_index * cols_per_row + col_index
#                     if img_index < len(page_imgs):
#                         img_path: Path = page_imgs[img_index]
#                         try:
#                             processed_path: Path = _process_image_for_doc(img_path, temp_dir)

#                             run = row_cells[col_index].paragraphs[0].add_run()
#                             run.add_picture(str(processed_path), width=img_width)

#                             caption = row_cells[col_index].add_paragraph(img_path.stem)
#                             caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#                             caption.runs[0].font.name = "Aptos"
#                             caption.runs[0].font.size = Pt(10)

#                         except Exception as e:
#                             print(f"Skipping {img_path.name} in document due to an error: {e}")

#             if start + images_per_page < len(image_paths):
#                 doc.add_page_break()

#         output_path: Path = save_directory / "Photo_Appendix.docx"
#         doc.save(output_path)
#         print(f"\nWord document successfully saved to: {output_path}")
#         messagebox.showinfo("Success", f"Word document created at:\n{output_path}")

#     finally:
#         shutil.rmtree(temp_dir)
#         print("Cleaned up temporary directory.")


# # =====================================================================
# # Controller: Main orchestration logic
# # =====================================================================

# def main_orchestrator() -> None:
#     """
#     Coordinates the end-to-end flow:
#     - Ask user for a source folder of images.
#     - Optionally load an Excel list to filter which images to include.
#     - Optionally compress selected images into a safe subfolder.
#     - Build the Word photo appendix from the selected set.

#     Rationale for GUI prompts:
#     - This script is designed for non-technical users who prefer dialogs.
#     - Each decision point is explicit to avoid accidental destructive actions.
#     """
#     root = tk.Tk()
#     root.withdraw()

#     source_dir_str: str = filedialog.askdirectory(title="Select the folder of photos for the report")
#     if not source_dir_str:
#         print("No source folder selected. Exiting.")
#         return
#     source_dir: Path = Path(source_dir_str)

#     all_images: List[Path] = [
#         p for p in source_dir.glob("*")
#         if p.is_file() and p.suffix.lower() in {".jpg", ".jpeg", ".png", ".bmp"}
#     ]

#     photos_for_report: List[Path] = []
#     sort_for_doc: bool = True

#     if messagebox.askyesno("Filter Photos?", "Do you want to use an Excel file to select specific photos for the report?"):
#         required_photo_list: List[str] = get_photo_list_from_excel()
#         if not required_photo_list:
#             print("No photo list retrieved from Excel, or dialog was cancelled. Aborting process.")
#             return

#         sort_numeric: bool = messagebox.askyesno(
#             "Photo ordering",
#             "How should photos be ordered in the Word document?\n\n"
#             "Yes = sort numerically (ascending)\n"
#             "No = use the native Excel order"
#         )

#         ordered_stems: List[str] = (
#             sorted(required_photo_list, key=natural_sort_key_string) if sort_numeric else required_photo_list
#         )

#         # When Excel drives ordering, we must not re-sort during Word generation.
#         sort_for_doc = False

#         # Using a lookup avoids repeatedly scanning the folder list, which matters on large directories.
#         image_by_stem: Dict[str, Path] = {p.stem: p for p in all_images}

#         missing_stems: List[str] = []
#         for stem in ordered_stems:
#             stem_clean: str = str(stem).strip()
#             match: Optional[Path] = image_by_stem.get(stem_clean)
#             if match is not None:
#                 photos_for_report.append(match)
#             else:
#                 missing_stems.append(stem_clean)

#         print(f"\nMatched {len(photos_for_report)} photo(s) from the selected folder.")

#         if missing_stems:
#             print(f"Missing {len(missing_stems)} photo(s) listed in Excel (not found in folder).")
#             messagebox.showwarning(
#                 "Some photos not found",
#                 f"{len(missing_stems)} photo(s) listed in Excel were not found in the selected folder.\n\n"
#                 "The document will be created with the photos that were found."
#             )

#     else:
#         photos_for_report = all_images
#         sort_for_doc = True

#     if not photos_for_report:
#         print("No matching photos were found for the report.")
#         messagebox.showwarning("No Photos", "No photos were found matching the criteria.")
#         return

#     print(f"\nFound {len(photos_for_report)} photos for the report.")

#     if messagebox.askyesno(
#         "Confirm Compression",
#         f"Do you want to save compressed copies of the {len(photos_for_report)} selected images?\n\n"
#         "This will create a new 'COMPRESSED_IMAGES' folder. Original files will NOT be changed."
#     ):
#         quality: Optional[int] = simpledialog.askinteger(
#             "Compression Quality",
#             "Enter JPEG quality for compression (1-100, higher is better quality):",
#             initialvalue=50, minvalue=1, maxvalue=100
#         )
#         effective_quality: int = 50 if quality is None else int(quality)
#         save_compressed_images(photos_for_report, source_dir, quality=effective_quality)

#     if messagebox.askyesno("Confirm Document Creation", f"Create a Word document with the {len(photos_for_report)} selected photos?"):
#         create_photo_document(photos_for_report, source_dir, sort_images=sort_for_doc)

#     print("\nProcess finished.")


# # =====================================================================
# # Script entry point
# # =====================================================================

# if __name__ == "__main__":
#     main_orchestrator()



# # # =====================================================================
# # # Photo appendix builder
# # # =====================================================================
# # # Purpose:
# # # - Read an optional Excel list of required photo base names.
# # # - Let the user choose a source folder of images.
# # # - Optionally compress copies of the selected images in a safe subfolder.
# # # - Generate a Word document in landscape layout with a 2x3 photo grid per page.
# # #
# # # Teaching notes:
# # # - GUI interactions use Tkinter's modal dialogs (no full window UI).
# # # - Images are handled with Pillow (rotation by EXIF, RGB conversion).
# # # - Word doc generation uses python-docx tables to implement the grid layout.
# # # - Excel reading uses openpyxl in read-only mode for robustness.
# # #
# # # Style choices for you:
# # # - Australian English spelling in comments and messages.
# # # - No ordered-number comments inside code blocks.
# # # - Type hints everywhere to make intent explicit.
# # # - Comments focus on why something is done, not only what.
# # # =====================================================================

# # # --- Module imports ---

# # import tkinter as tk
# # from tkinter import filedialog, messagebox
# # from pathlib import Path
# # import shutil
# # import tempfile
# # from typing import List, Optional

# # from PIL import Image
# # from docx import Document
# # from docx.shared import Inches, Pt
# # from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# # from openpyxl import load_workbook
# # import re


# # # =====================================================================
# # # Module 1: Data input and image processing
# # # =====================================================================

# # def get_photo_list_from_excel() -> List[str]:
# #     """
# #     Prompts the user for an Excel file and extracts values from the first column
# #     of the active worksheet. Each non-empty cell is treated as a required photo
# #     base name (matching Path.stem) to filter the source images.

# #     Why this design:
# #     - Using read_only=True avoids loading the entire workbook into memory,
# #       which helps on large files and is safer for simple read tasks.
# #     - Reading cells row-by-row avoids edge cases with high-level slicing.
# #     """
# #     print("\nPlease select the Excel file containing the list of required photo filenames.")
# #     file_path_str: str = filedialog.askopenfilename(
# #         title="Select Excel Photo List (.xlsx)",
# #         filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
# #     )

# #     if not file_path_str:
# #         print("No Excel file selected.")
# #         return []

# #     file_path: Path = Path(file_path_str)

# #     try:
# #         workbook = load_workbook(filename=file_path, read_only=True)
# #         sheet = workbook.active

# #         photo_list: List[str] = []
# #         # Iterating from 1..max_row matches openpyxl’s 1-based indexing.
# #         for i in range(1, sheet.max_row + 1):
# #             cell = sheet.cell(row=i, column=1)
# #             if cell.value is not None:
# #                 # Casting to str ensures uniform matching against Path.stem later.
# #                 photo_list.append(str(cell.value))

# #         print(f"Successfully read {len(photo_list)} entries from Excel file.")
# #         return photo_list

# #     except Exception as e:
# #         # User feedback is important for diagnosing common issues such as locked files,
# #         # sheet protection, or the wrong file format.
# #         print(f"Error reading Excel file: {e}")
# #         messagebox.showerror("File Error", f"Could not read the selected Excel file.\n\nError: {e}")
# #         return []


# # def _get_compressed_image(img_path: Path) -> Image.Image:
# #     """
# #     Loads an image with Pillow and ensures it is in a JPEG-compatible mode.

# #     Why convert modes:
# #     - JPEG does not support alpha channels or palette modes.
# #     - Converting to RGB avoids runtime errors when saving as JPEG.
# #     """
# #     img: Image.Image = Image.open(img_path)
# #     if img.mode in ("RGBA", "P"):
# #         img = img.convert("RGB")
# #     return img


# # def save_compressed_images(image_paths: List[Path], source_directory: Path, quality: int = 75) -> None:
# #     """
# #     Writes compressed copies of the given images into 'COMPRESSED_IMAGES' in the source folder.

# #     Non-destructive by design:
# #     - Originals are never overwritten.
# #     - Copies go into a dedicated subfolder for easy review and rollback.

# #     About JPEG quality:
# #     - Must be in the 1-100 range.
# #     - 75 is a reasonable default that saves space while maintaining visible quality.
# #     """
# #     if not image_paths:
# #         print("No images were provided for compression.")
# #         return

# #     output_dir: Path = source_directory / "COMPRESSED_IMAGES"
# #     output_dir.mkdir(exist_ok=True)

# #     print(f"\n--- Starting image compression (saving to '{output_dir.name}') ---")
# #     compressed_count: int = 0
# #     skipped_count: int = 0

# #     for img_path in image_paths:
# #         try:
# #             # File size fetch allows a simple before/after compression report.
# #             original_size_kb: float = img_path.stat().st_size / 1024.0

# #             # Preparing the image ensures valid JPEG mode.
# #             compressed_img = _get_compressed_image(img_path)
# #             save_path: Path = output_dir / img_path.name

# #             # optimise=True lets Pillow attempt Huffman-table and DCT optimisations.
# #             compressed_img.save(save_path, format="JPEG", quality=quality, optimize=True)

# #             new_size_kb: float = save_path.stat().st_size / 1024.0
# #             reduction_kb: float = max(0.0, original_size_kb - new_size_kb)
# #             reduction_pct: float = (reduction_kb / original_size_kb * 100.0) if original_size_kb > 0 else 0.0

# #             print(f"-> Saved compressed: {img_path.name} ({original_size_kb:.1f} KB -> {new_size_kb:.1f} KB, {reduction_pct:.1f}% smaller)")
# #             compressed_count += 1

# #         except Exception as e:
# #             # If a single file fails (corrupt image, permission issue), we log and continue.
# #             print(f"-> Error compressing {img_path.name}: {e}. Skipping.")
# #             skipped_count += 1

# #     print("\n--- Compression summary ---")
# #     print(f"Successfully saved {compressed_count} compressed file(s) to '{output_dir.name}'.")
# #     if skipped_count > 0:
# #         print(f"Skipped or failed: {skipped_count} file(s).")
# #     print("---------------------------")


# # def _process_image_for_doc(img_path: Path, temp_dir: Path) -> Path:
# #     """
# #     Produces a temporary JPEG tailored for Word insertion.

# #     Operations:
# #     - Opens the image.
# #     - Applies EXIF-based orientation fix if present.
# #     - Converts to RGB if required.
# #     - Saves to a temp JPG (quality ~85) and returns that temp path.

# #     Why use temp files:
# #     - Word embedding expects a stable on-disk file path.
# #     - Avoids any mutation of original images.
# #     """
# #     img: Image.Image = Image.open(img_path)

# #     # EXIF orientation handling: cameras often store rotation in metadata
# #     # rather than rotating the pixel matrix. This aligns the visual result.
# #     try:
# #         exif = getattr(img, "_getexif", lambda: None)()
# #         if exif:
# #             orientation_tag = 274  # standard orientation key
# #             orientation = exif.get(orientation_tag)
# #             if orientation == 3:
# #                 img = img.rotate(180, expand=True)
# #             elif orientation == 6:
# #                 img = img.rotate(270, expand=True)
# #             elif orientation == 8:
# #                 img = img.rotate(90, expand=True)
# #     except Exception:
# #         # If EXIF is malformed or absent, we simply proceed without rotation.
# #         pass

# #     if img.mode in ("RGBA", "P"):
# #         img = img.convert("RGB")

# #     temp_jpg_path: Path = temp_dir / f"{img_path.stem}_processed.jpg"
# #     img.save(temp_jpg_path, format="JPEG", quality=85)
# #     return temp_jpg_path

# # def natural_sort_key(path: Path) -> list:
# #     """
# #     Returns a list that mixes strings and integers to achieve natural (human-like) sorting.
# #     Example:
# #         DSCN3801 -> ['DSCN', 3801]
# #         10 -> [10]
# #     """
# #     # Convert the filename stem (without extension) into alternating text and number parts
# #     parts = re.split(r'(\d+)', path.stem)
# #     # Convert numeric substrings to integers for proper comparison
# #     return [int(p) if p.isdigit() else p for p in parts]

# # def create_photo_document(image_paths: List[Path], save_directory: Path) -> None:
# #     """
# #     Creates a landscape Word document with a 2x3 image grid per page.
# #     Each image is centred in its cell and labelled with its base filename.

# #     Layout decisions:
# #     - Using tables provides a robust grid in Word, which keeps images aligned.
# #     - Fixed image width lets Word keep aspect ratio and scale height automatically.
# #     - Page breaks are inserted after each complete page batch, except the last.
# #     """
# #     if not image_paths:
# #         messagebox.showinfo("No Images Found", "No images were selected or found to create the document.")
# #         return

# #     doc: Document = Document()
# #     section = doc.sections[-1]

# #     # Orientation value 1 corresponds to landscape for python-docx.
# #     # Swapping width and height is necessary to apply the change cleanly.
# #     section.orientation = 1
# #     section.page_width, section.page_height = section.page_height, section.page_width

# #     # Grid parameters define the visual structure.
# #     images_per_page: int = 6
# #     cols_per_row: int = 3
# #     img_width: Inches = Inches(2.5)

# #     # A temporary folder is used to store EXIF-fixed JPEGs for embedding.
# #     temp_dir: Path = Path(tempfile.mkdtemp())
# #     print(f"Created temporary directory for image processing: {temp_dir}")

# #     try:
# #         # Sorting is alphabetical by default with Path objects. If numeric ordering
# #         # on stems is required (e.g., 2 before 10), a custom key would be used.
# #         image_paths.sort(key=natural_sort_key)


# #         # The batch loop slices the list into page-sized chunks to control pagination.
# #         for start in range(0, len(image_paths), images_per_page):
# #             page_imgs: List[Path] = image_paths[start:start + images_per_page]

# #             # Two rows per page since 6 images with 3 columns yields 2 rows.
# #             for row_index in range(2):
# #                 table = doc.add_table(rows=1, cols=cols_per_row)
# #                 row_cells = table.rows[0].cells

# #                 # The inner column loop calculates which image occupies each cell.
# #                 # Mapping formula places images left-to-right, top-to-bottom.
# #                 for col_index in range(cols_per_row):
# #                     img_index: int = row_index * cols_per_row + col_index
# #                     if img_index < len(page_imgs):
# #                         img_path = page_imgs[img_index]
# #                         try:
# #                             processed_path = _process_image_for_doc(img_path, temp_dir)

# #                             # Each cell starts with a default empty paragraph. A run is the unit
# #                             # that holds either text or a picture. We attach the picture to a run.
# #                             run = row_cells[col_index].paragraphs[0].add_run()
# #                             run.add_picture(str(processed_path), width=img_width)

# #                             # Captions below images help with traceability in reports.
# #                             caption = row_cells[col_index].add_paragraph(img_path.stem)
# #                             caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# #                             caption.runs[0].font.name = "Aptos"
# #                             caption.runs[0].font.size = Pt(10)

# #                         except Exception as e:
# #                             # A single failure must not derail the document generation.
# #                             print(f"Skipping {img_path.name} in document due to an error: {e}")

# #             # Adding a page break after a full page keeps paging predictable.
# #             # The final page does not require an extra break at the end.
# #             if start + images_per_page < len(image_paths):
# #                 doc.add_page_break()

# #         output_path: Path = save_directory / "Photo_Appendix.docx"
# #         doc.save(output_path)
# #         print(f"\nWord document successfully saved to: {output_path}")
# #         messagebox.showinfo("Success", f"Word document created at:\n{output_path}")

# #     finally:
# #         # Temporary artefacts are removed to keep the filesystem tidy.
# #         shutil.rmtree(temp_dir)
# #         print("Cleaned up temporary directory.")


# # # =====================================================================
# # # Controller: Main orchestration logic
# # # =====================================================================

# # def main_orchestrator() -> None:
# #     """
# #     Coordinates the end-to-end flow:
# #     - Ask user for a source folder of images.
# #     - Optionally load an Excel list to filter which images to include.
# #     - Optionally compress selected images into a safe subfolder.
# #     - Build the Word photo appendix from the selected set.

# #     Rationale for GUI prompts:
# #     - This script is designed for non-technical users who prefer dialogs.
# #     - Each decision point is explicit to avoid accidental destructive actions.
# #     """
# #     # The hidden root window allows modal dialogs without a visible empty frame.
# #     root = tk.Tk()
# #     root.withdraw()

# #     source_dir_str: str = filedialog.askdirectory(title="Select the folder of photos for the report")
# #     if not source_dir_str:
# #         print("No source folder selected. Exiting.")
# #         return
# #     source_dir: Path = Path(source_dir_str)

# #     # Collecting images by extension keeps behaviour simple and predictable.
# #     # Using glob('*') plus suffix filtering allows arbitrary file names.
# #     all_images: List[Path] = [
# #         p for p in source_dir.glob('*')
# #         if p.is_file() and p.suffix.lower() in {'.jpg', '.jpeg', '.png', '.bmp'}
# #     ]

# #     # This list will contain the final set of images to process downstream.
# #     photos_for_report: List[Path] = []

# #     # Users can choose to limit the set to an Excel-provided list of base names.
# #     # This is handy when a report requires a specific subset from a large folder.
# #     if messagebox.askyesno("Filter Photos?", "Do you want to use an Excel file to select specific photos for the report?"):
# #         required_photo_list: List[str] = get_photo_list_from_excel()
# #         if not required_photo_list:
# #             print("No photo list retrieved from Excel, or dialog was cancelled. Aborting process.")
# #             return

# #         required_set = set(required_photo_list)
# #         print(f"\nLoaded {len(required_set)} required photo filenames from Excel.")

# #         # Matching on Path.stem lets Excel hold base names without extensions.
# #         for image in all_images:
# #             if image.stem in required_set:
# #                 photos_for_report.append(image)
# #     else:
# #         # If no filter is requested, everything in the folder is included.
# #         photos_for_report = all_images

# #     if not photos_for_report:
# #         print("No matching photos were found for the report.")
# #         messagebox.showwarning("No Photos", "No photos were found matching the criteria.")
# #         return

# #     print(f"\nFound {len(photos_for_report)} photos for the report.")

# #     # Compression is offered as a separate opt-in because it is not required
# #     # for the Word document and can take additional time on large sets.
# #     if messagebox.askyesno(
# #         "Confirm Compression",
# #         f"Do you want to save compressed copies of the {len(photos_for_report)} selected images?\n\n"
# #         "This will create a new 'COMPRESSED_IMAGES' folder. Original files will NOT be changed."
# #     ):
# #         # The quality prompt returns None if the user cancels; a fallback keeps the flow predictable.
# #         quality: Optional[int] = tk.simpledialog.askinteger(
# #             "Compression Quality",
# #             "Enter JPEG quality for compression (1-100, higher is better quality):",
# #             initialvalue=50, minvalue=1, maxvalue=100
# #         )
# #         effective_quality: int = 50 if quality is None else int(quality)
# #         save_compressed_images(photos_for_report, source_dir, quality=effective_quality)

# #     # The document generation is also explicit to give users a last chance to bail out.
# #     if messagebox.askyesno("Confirm Document Creation", f"Create a Word document with the {len(photos_for_report)} selected photos?"):
# #         create_photo_document(photos_for_report, source_dir)

# #     print("\nProcess finished.")


# # # =====================================================================
# # # Script entry point
# # # =====================================================================

# # if __name__ == "__main__":
# #     main_orchestrator()


